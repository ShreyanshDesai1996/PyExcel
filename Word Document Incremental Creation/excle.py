import docx


def setDateInFile(date, original, file):
    doc = docx.Document(original)
    # all_paras = doc.paragraphs
    for paragraph in doc.paragraphs:
        if "{{}}" in paragraph.text:
            # print(paragraph.text)
            # paragraph.text = paragraph.text.replace("{{}}", date)
            # print("Date set to" + date)
            inline = paragraph.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if "{{}}" in inline[i].text:
                    text = inline[i].text.replace("{{}}", date)
                    inline[i].text = text
            print(paragraph.text)
    doc.save(file)


dates = [
    "12/01/2022",
    "13/01/2022",
    "14/01/2022",
    "15/01/2022",
    "16/01/2022",
    "17/01/2022",
    "18/01/2022",
    "19/01/2022",
    "20/01/2022",
    "21/01/2022",
    "22/01/2022",
    "23/01/2022",
    "24/01/2022",
    "25/01/2022",
    "26/01/2022",
    "27/01/2022",
    "28/01/2022",
    "29/01/2022",
    "30/01/2022",
    "31/01/2022",
    "01/02/2022",
    "02/02/2022",
    "03/02/2022",
    "04/02/2022",
    "05/02/2022",
    "06/02/2022",
    "07/02/2022",
    "08/02/2022",
    "09/02/2022",
    "10/02/2022",
    "11/02/2022",
    "12/02/2022",
    "13/02/2022",
    "14/02/2022",
    "15/02/2022",
]


for date in dates:
    # mydoc = docx.Document()
    # src = "C:/Users/Shrey/Downloads/R9073554_PRAMILA_DESAI_161221223347.docx"
    # dest = "C:/Users/Shrey/Downloads/" + date.replace("/", "_") + ".docx"
    # path = shutil.copyfile(src, dest)
    setDateInFile(
        date,
        "C:/Users/Shrey/Downloads/R9073554_PRAMILA_DESAI_161221223347.docx",
        "C:/Users/Shrey/Downloads/" + date.replace("/", "_") + ".docx",
    )
